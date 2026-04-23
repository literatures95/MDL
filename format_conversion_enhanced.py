"""
MDL 格式转换增强模块 - 改进实施代码框架
快速参考: 直接替换或补充 formats.py 中的实现

优先级实施顺序:
P0 (立即): PDF API + Excel 增强 + URL 支持
P1 (中期): HTML 改进 + DOCX 图片
P2 (优化): 并行处理 + PPTX 增强
"""

# =============================================================================
# P0 优先级: PDF 转换增强 (pdfplumber 本地方案)
# =============================================================================

def pdf_to_markdown_pdfplumber(pdf_path: str) -> str:
    """
    使用 pdfplumber 提升 PDF 转换质量
    安装: pip install pdfplumber
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("需要安装 pdfplumber: pip install pdfplumber")
    
    markdown_parts = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            markdown_parts.append(f"## 第 {page_idx + 1} 页\n")
            
            # 1. 提取文本
            text = page.extract_text()
            if text:
                markdown_parts.append(text)
            
            # 2. 提取表格 (pdfplumber 专长)
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    markdown_parts.append("\n")
                    markdown_parts.append(table_to_markdown(table))
            
            # 3. 提取图片 (可选)
            for img in page.images:
                markdown_parts.append(f"\n![图片]({img['name']})\n")
            
            markdown_parts.append("\n---\n")
    
    return "\n".join(markdown_parts)


def table_to_markdown(table: list) -> str:
    """将 pdfplumber 表格转换为 Markdown"""
    if not table or not table[0]:
        return ""
    
    # 确定列数
    max_cols = max(len(row) for row in table)
    
    # 规范化：填充空行
    normalized = []
    for row in table:
        normalized_row = row + [None] * (max_cols - len(row))
        normalized.append([str(cell) if cell else "" for cell in normalized_row])
    
    # 生成 Markdown
    header = "| " + " | ".join(normalized[0]) + " |"
    separator = "| " + " | ".join(["---"] * max_cols) + " |"
    body = [
        "| " + " | ".join(row) + " |"
        for row in normalized[1:]
    ]
    
    return header + "\n" + separator + "\n" + "\n".join(body)


# P0 选项 B: 使用 TextIn API (付费，但质量最佳)
# ===========================================================================

def pdf_to_markdown_textin_api(pdf_path: str, app_id: str, secret_code: str) -> str:
    """
    使用 TextIn API 进行高质量 PDF 转换
    
    官网: https://www.textin.com/
    注册获取 API 密钥
    
    成本: ~0.05 元/页
    """
    import requests
    
    url = "https://api.textin.com/ai/service/v1/pdf_to_markdown"
    
    headers = {
        "x-ti-app-id": app_id,
        "x-ti-secret-code": secret_code
    }
    
    params = {
        "apply_document_tree": 1,      # 保留文档结构
        "markdown_details": 1,         # 增强元数据
        "table_flavor": "md",          # 表格格式为 Markdown
        "get_image": "objects",        # 提取图片
        "dpi": 144
    }
    
    with open(pdf_path, "rb") as f:
        response = requests.post(url, headers=headers, params=params, data=f)
    
    result = response.json()
    
    if result.get("code") == 200:
        return result["result"]["markdown"]
    else:
        raise Exception(f"TextIn API 错误: {result.get('message')}")


# =============================================================================
# P0 优先级: Excel 转换增强
# =============================================================================

def xlsx_to_markdown_enhanced(xlsx_path: str, sheet_names: list = None,
                             detect_header: bool = True) -> str:
    """
    增强的 Excel 转换器
    
    参数:
        sheet_names: 指定转换的工作表列表，None=所有
        detect_header: 是否自动检测表头行
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("需要安装 openpyxl: pip install openpyxl")
    
    workbook = load_workbook(xlsx_path)
    
    # 工作表选择
    if sheet_names:
        sheets_to_process = [
            workbook[name] for name in sheet_names
            if name in workbook.sheetnames
        ]
    else:
        sheets_to_process = workbook.worksheets
    
    markdown_parts = []
    
    for sheet in sheets_to_process:
        markdown_parts.append(f"## {sheet.title}\n")
        
        # 自动检测表头行 (启发式)
        header_row_idx = 0
        if detect_header:
            for idx in range(min(5, sheet.max_row)):
                row = list(sheet.iter_rows(min_row=idx+1, max_row=idx+1, values_only=True))[0]
                # 如果行的所有值都非空，可能是表头
                if all(cell is not None for cell in row):
                    header_row_idx = idx
                    break
        
        # 生成表格
        first_row = True
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            cells = [str(cell) if cell is not None else "" for cell in row]
            
            if any(cells):  # 跳过空行
                md_row = "| " + " | ".join(cells) + " |"
                markdown_parts.append(md_row)
                
                # 在检测到的表头下添加分隔符
                if first_row and row_idx >= header_row_idx:
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    markdown_parts.append(separator)
                    first_row = False
        
        markdown_parts.append("\n")
    
    return "\n".join(markdown_parts)


def get_xlsx_sheet_names(xlsx_path: str) -> list:
    """获取 Excel 文件的所有工作表名称"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("需要安装 openpyxl")
    
    workbook = load_workbook(xlsx_path)
    return workbook.sheetnames


# =============================================================================
# P0 优先级: URL 网页处理
# =============================================================================

def url_to_markdown(url: str) -> str:
    """
    将网页 URL 转换为 Markdown
    """
    import requests
    from bs4 import BeautifulSoup
    
    try:
        response = requests.get(url, timeout=10)
        response.encoding = 'utf-8'
        html = response.text
    except requests.RequestException as e:
        raise Exception(f"无法获取网页: {e}")
    
    # 使用 html2text 进行转换
    try:
        from html2text import HTML2Text
        h = HTML2Text()
        h.body_width = 0  # 不换行
        markdown = h.handle(html)
        return markdown
    except ImportError:
        # 回退方案：BeautifulSoup + 正则表达式
        soup = BeautifulSoup(html, 'html.parser')
        
        # 移除脚本和样式
        for tag in soup(['script', 'style']):
            tag.decompose()
        
        # 简单的 BeautifulSoup 转换
        return soup.get_text()


def extract_links_from_url(url: str) -> list:
    """
    从网页 URL 提取所有链接
    """
    import requests
    from bs4 import BeautifulSoup
    from urllib.parse import urljoin, urlparse
    
    response = requests.get(url, timeout=10)
    response.encoding = 'utf-8'
    
    soup = BeautifulSoup(response.text, 'html.parser')
    
    links = []
    base_domain = urlparse(url).netloc
    
    for a_tag in soup.find_all('a', href=True):
        href = a_tag['href']
        full_url = urljoin(url, href)
        
        # 只保留同域名的链接
        if urlparse(full_url).netloc == base_domain:
            links.append(full_url)
    
    return list(set(links))  # 去重


def batch_convert_urls(urls: list, output_dir: str = "converted_urls") -> dict:
    """
    批量转换多个 URL 为 Markdown 文件
    
    返回: {url: {status, file_path, error}}
    """
    import os
    from urllib.parse import urlparse
    
    os.makedirs(output_dir, exist_ok=True)
    results = {}
    
    for url in urls:
        try:
            # 获取网页内容
            markdown = url_to_markdown(url)
            
            # 生成文件名
            path = urlparse(url).path
            filename = path.split('/')[-1] or 'index'
            if not filename.endswith('.md'):
                filename += '.md'
            
            file_path = os.path.join(output_dir, filename)
            
            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(markdown)
            
            results[url] = {
                'status': 'success',
                'file': file_path
            }
        
        except Exception as e:
            results[url] = {
                'status': 'error',
                'error': str(e)
            }
    
    return results


# =============================================================================
# P1 优先级: HTML 转换增强
# =============================================================================

def html_to_markdown_beautifulsoup(html: str, remove_nav: bool = True) -> str:
    """
    使用 BeautifulSoup 转换 HTML (比正则表达式更可靠)
    安装: pip install beautifulsoup4 html2text
    """
    from bs4 import BeautifulSoup
    
    soup = BeautifulSoup(html, 'html.parser')
    
    # 移除导航和脚本
    if remove_nav:
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            tag.decompose()
    
    # 使用 html2text 进行高质量转换
    try:
        from html2text import HTML2Text
        h = HTML2Text()
        h.body_width = 0
        h.protect_links = True
        return h.handle(str(soup))
    except ImportError:
        # 回退: 使用简单的文本提取
        return soup.get_text()


def extract_links_from_html(html: str, base_url: str = "") -> list:
    """
    从 HTML 提取所有链接
    """
    from bs4 import BeautifulSoup
    from urllib.parse import urljoin
    
    soup = BeautifulSoup(html, 'html.parser')
    links = []
    
    for a_tag in soup.find_all('a', href=True):
        href = a_tag['href']
        if base_url:
            href = urljoin(base_url, href)
        links.append(href)
    
    return links


# =============================================================================
# P1 优先级: DOCX 图片提取
# =============================================================================

def docx_to_markdown_with_images(docx_path: str, extract_images: bool = True) -> str:
    """
    增强的 DOCX 转换器，支持图片提取
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    
    import os
    from pathlib import Path
    
    doc = Document(docx_path)
    lines = []
    
    # 创建图片目录
    image_dir = None
    if extract_images:
        docx_stem = Path(docx_path).stem
        image_dir = f"{docx_stem}_images"
        os.makedirs(image_dir, exist_ok=True)
    
    # 处理段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 处理标题
        style_name = para.style.name.lower()
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        else:
            # 处理格式
            md_text = ""
            for run in para.runs:
                segment = run.text
                if run.bold and run.italic:
                    segment = f"***{segment}***"
                elif run.bold:
                    segment = f"**{segment}**"
                elif run.italic:
                    segment = f"*{segment}*"
                md_text += segment
            lines.append(md_text)
    
    # 处理表格
    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")
    
    # 提取图片
    if extract_images and image_dir:
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image = rel.target_part.blob
                    image_filename = f"image_{image_count}.png"
                    image_path = os.path.join(image_dir, image_filename)
                    with open(image_path, 'wb') as f:
                        f.write(image)
                    lines.append(f"![Image]({image_dir}/{image_filename})")
                    image_count += 1
                except Exception:
                    pass
    
    return "\n".join(lines)


# =============================================================================
# 集成到 MDL 的使用示例
# =============================================================================

if __name__ == "__main__":
    # P0 示例
    
    # 1. PDF 转换
    # pdf_md = pdf_to_markdown_pdfplumber("document.pdf")
    
    # 2. Excel 转换
    # sheets = get_xlsx_sheet_names("data.xlsx")
    # excel_md = xlsx_to_markdown_enhanced("data.xlsx", sheet_names=sheets[:2])
    
    # 3. URL 转换
    # url_md = url_to_markdown("https://example.com")
    # links = extract_links_from_url("https://example.com")
    # batch_convert_urls(["https://example.com/page1", "https://example.com/page2"])
    
    print("MDL 格式转换增强模块")
    print("优先级 P0: PDF (pdfplumber) + Excel + URL")
    print("优先级 P1: HTML (BeautifulSoup) + DOCX 图片")
    print("优先级 P2: 并行处理、PPTX 增强")
